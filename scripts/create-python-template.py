from docx import Document
from docx.shared import Pt, Inches
import json

doc = Document()

# Add a title
doc.add_heading('Cost Analysis Report', 0)

# Add a paragraph with a merge field
p = doc.add_paragraph()
p.add_run('Application: ')
# Add a simple text placeholder
p.add_run('{{appName}}')

# Add another paragraph
p2 = doc.add_paragraph()
p2.add_run('Generated: ')
p2.add_run('{{generatedAt}}')

# Add a summary section
doc.add_heading('Summary', 1)
doc.add_paragraph('{{summary}}')

# Add a services table
doc.add_heading('Identified Services', 1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Provider'
hdr[2].text = 'Category'
hdr[3].text = 'Monthly Cost'

# Add a row with placeholder
row = table.add_row().cells
row[0].text = '{{serviceName}}'
row[1].text = '{{serviceProvider}}'
row[2].text = '{{serviceCategory}}'
row[3].text = '{{serviceCost}}'

# Save
doc.save('templates/python-docx-template.docx')
print('Template saved successfully')
