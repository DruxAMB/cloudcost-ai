"""
Create a DOCX template with Word content controls (SDT elements) that
Doctavian's Maven Mule engine might recognize.

The Maven Mule Template Builder add-in inserts content controls with
specific tags. We'll create SDT elements with tags matching our data fields.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Add a title
doc.add_heading('Cost Analysis Report', 0)

def add_content_control(paragraph, tag_name, placeholder_text):
    """Add a Word content control (SDT) with a specific tag to a paragraph."""
    # Create the SDT element
    sdt = OxmlElement('w:sdt')
    
    # SDT properties
    sdtPr = OxmlElement('w:sdtPr')
    tag = OxmlElement('w:tag')
    tag.set(qn('w:val'), tag_name)
    sdtPr.append(tag)
    
    id_elem = OxmlElement('w:id')
    id_elem.set(qn('w:val'), str(hash(tag_name) % 100000))
    sdtPr.append(id_elem)
    
    # Add a text element to indicate it's a plain text content control
    text_elem = OxmlElement('w:text')
    sdtPr.append(text_elem)
    
    sdt.append(sdtPr)
    
    # SDT content
    sdtContent = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = placeholder_text
    run.append(text)
    sdtContent.append(run)
    
    sdt.append(sdtContent)
    
    paragraph._p.append(sdt)
    return sdt

# Add paragraph with content control for appName
p1 = doc.add_paragraph()
p1.add_run('Application: ')
add_content_control(p1, 'appName', 'App Name Placeholder')

# Add paragraph with content control for generatedAt
p2 = doc.add_paragraph()
p2.add_run('Generated: ')
add_content_control(p2, 'generatedAt', 'Date Placeholder')

# Add summary section
doc.add_heading('Summary', 1)
p3 = doc.add_paragraph()
add_content_control(p3, 'summary', 'Summary Placeholder')

# Add a services table
doc.add_heading('Identified Services', 1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Provider'
hdr[2].text = 'Category'
hdr[3].text = 'Monthly Cost'

# Save
doc.save('templates/sdt-content-control-template.docx')
print('SDT content control template saved successfully')

# Also create a version with {{}} style placeholders inside the SDT
doc2 = Document()
doc2.add_heading('Cost Analysis Report', 0)

# Simple text with mustache-style placeholders
p = doc2.add_paragraph('Application: {{appName}}')
p = doc2.add_paragraph('Generated: {{generatedAt}}')
doc2.add_heading('Summary', 1)
doc2.add_paragraph('{{summary}}')

doc2.save('templates/mustache-template.docx')
print('Mustache template saved successfully')
